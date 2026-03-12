# coding: utf-8
"""pages/04_cam.py — Module 4: CAM Generator (8 Steps)"""

import json, os, io
from datetime import datetime, date

import matplotlib
matplotlib.use("Agg")
import matplotlib.pyplot as plt
import matplotlib.patches as mpatches
import pandas as pd
import streamlit as st

from utils.gemini_client import call_gemini_with_retry
from utils.ui_icons import svg_header, get_svg, icon_label

# ── Page config ───────────────────────────────────────────────────────────────
st.set_page_config(page_title="Module 4 — CAM Generator", layout="wide")
st.markdown("""
<style>
@import url('https://fonts.googleapis.com/css2?family=Inter:wght@400;600;700;800&display=swap');
html, body, [class*="css"] { font-family: 'Inter', sans-serif; }
.stApp { background: #f8fafc; }
h1,h2,h3,h4 { color: #1B3A6B; }
.stButton > button { background:#1B3A6B!important;color:white!important;border-radius:7px!important;font-weight:600!important; }
.stButton > button:hover { background:#2563EB!important; }
.cam-section-card { background:white;border-radius:10px;padding:1rem 1.2rem;
    box-shadow:0 2px 6px rgba(0,0,0,0.07);border-left:4px solid #1B3A6B;margin-bottom:.6rem; }
.conf-high   { color:#059669;font-weight:700; }
.conf-med    { color:#d97706;font-weight:700; }
.conf-low    { color:#dc2626;font-weight:700; }
.india-row   { background:#fff7ed;border-left:3px solid #ea580c;border-radius:5px;
               padding:.4rem .8rem;margin:.25rem 0;font-size:.85rem;color:#9a3412; }
</style>
""", unsafe_allow_html=True)

os.makedirs("data", exist_ok=True)

# ── STEP 1: Session persistence ───────────────────────────────────────────────
def load_payload(name: str, filepath: str, required_step: int):
    if name not in st.session_state:
        try:
            with open(filepath) as f:
                st.session_state[name] = json.load(f)
            st.info(f"Session restored from {filepath}")
        except FileNotFoundError:
            _page_map = {1: "pages/01_ingestor.py", 2: "pages/02_research.py", 3: "pages/03_recommendation.py"}
            st.warning(f"{name} not found. Please complete Step {required_step} first.", icon=":material/warning:")
            if st.button(f"← Go to Step {required_step}"):
                st.switch_page(_page_map[required_step])
            st.stop()

load_payload("extraction_payload",     "./data/extraction_payload.json",     1)
load_payload("research_payload",       "./data/research_payload.json",       2)
load_payload("recommendation_payload", "./data/recommendation_payload.json", 3)

extraction_payload = st.session_state["extraction_payload"]
research_payload   = st.session_state["research_payload"]
recommendation_payload = st.session_state["recommendation_payload"]
EP = extraction_payload # keep for back compat if missed anywhere
RP = research_payload
REC = recommendation_payload

entity     = EP.get("entity_context", {}) or {}
financials = EP.get("financials", {})      or {}
doc_flags  = EP.get("fraud_flags", [])     or []
tri_flags  = RP.get("triangulation_flags", []) or []
ro         = RP.get("research_output", {}) or {}
pi         = RP.get("primary_insights", {}) or {}
ewi        = ro.get("early_warning_signals", []) or []

company_name = entity.get("company_name", "Entity")
sector       = entity.get("sector", "Other")
loan_amount  = float(entity.get("loan_amount_cr") or entity.get("loan_amount") or 0)
cin          = entity.get("cin", "N/A")
decision     = REC.get("decision", "N/A")
max_loan     = REC.get("recommended_loan_cr", 0)
rate         = REC.get("recommended_rate_pct", 0)
pd_val       = REC.get("pd", 0)
lgd_val      = REC.get("lgd", 0)
z_score      = REC.get("z_score", 0)
confidence   = REC.get("confidence", 0)
india_concerns = REC.get("india_specific_concerns", []) or []

# ── Page header ───────────────────────────────────────────────────────────────
svg_header("REPORT", "Module 4 — Credit Appraisal Memo", level=1)
st.caption(f"**{company_name}** | {sector} | Loan: ₹{loan_amount:.0f}Cr | Decision: **{decision}**")

# ── Step progress bar ─────────────────────────────────────────────────────────
steps_done = sum([
    "extraction_payload" in st.session_state,
    "research_payload"   in st.session_state,
    "recommendation_payload" in st.session_state,
])
st.progress(steps_done / 4, text=f"Pipeline: {steps_done}/4 modules complete")
st.divider()

# ── STEP 1: Data completeness check ──────────────────────────────────────────
REQUIRED_FIELDS = [
    ("financials.revenue_cr",                 financials.get("revenue_cr")),
    ("financials.ebitda_cr",                  financials.get("ebitda_cr")),
    ("financials.total_debt_cr",              financials.get("total_debt_cr")),
    ("financials.cibil_commercial_score",     financials.get("cibil_commercial_score")),
    ("financials.gst_2a_vs_3b_variance_pct",  financials.get("gst_2a_vs_3b_variance_pct")),
    ("recommendation_payload.decision",       REC.get("decision")),
    ("recommendation_payload.swot",           REC.get("swot")),
    ("recommendation_payload.india_specific_concerns", REC.get("india_specific_concerns")),
    ("research_payload.sector_risk_score",    ro.get("sector_risk_score")),
    ("research_payload.triangulation_flags",  tri_flags),
]
missing = [name for name, val in REQUIRED_FIELDS if not val and val != 0]
section_confidence_penalty = len(missing) * 0.15
if missing:
    with st.expander(f"{len(missing)} missing fields — some sections will have lower confidence"):
        for m in missing:
            st.warning(f"Missing: `{m}` — analyst review required", icon=":material/warning:")

# ── STEP 2: Multi-agent Gemini CAM generation ─────────────────────────────────
CAM_SECTIONS = [
    "committee_summary", "executive_summary", "company_overview",
    "promoter_background", "financial_analysis", "industry_outlook",
    "legal_risks", "five_cs_assessment", "swot_analysis",
    "key_risk_drivers", "early_warning_indicators", "final_recommendation",
]

def build_cam_prompt() -> str:
    cibil   = financials.get("cibil_commercial_score", "N/A")
    gst_var = financials.get("gst_2a_vs_3b_variance_pct", 0) or 0
    ecourt  = pi.get("ecourt_cases_found", 0) or 0
    mca_dt  = pi.get("mca_last_filing_date", "N/A")
    rbi_c   = pi.get("rbi_compliance", "N/A")
    swot    = json.dumps(REC.get("swot", {}), indent=2)
    five_cs = json.dumps(REC.get("five_cs", {}), indent=2)
    rationale = json.dumps(REC.get("decision_rationale", []))
    ewi_s   = json.dumps(ewi[:5], default=str)
    tri_s   = json.dumps(tri_flags[:5], default=str)
    flags_s = json.dumps(doc_flags[:5], default=str)

    # Fetch Z-band narrative directly
    from utils.credit_engine import CreditEngine
    engine = CreditEngine()
    z_band_label, z_band_narrative = engine.z_band(z_score)

    web_context = research_payload.get("web_context_used", {})
    def summarise_web_context(web_context: dict) -> str:
        lines = []
        search_map = {
            "news_context":     "General news search",
            "legal_context":    "Legal/litigation search (NCLT, court, default)",
            "sector_context":   "Sector outlook search",
            "promoter_context": "Promoter background search",
            "mca_context":      "MCA/ROC compliance search",
            "gst_context":      "GST compliance/notice search",
        }
        for key, label in search_map.items():
            content = web_context.get(key, "")
            if content and len(content.strip()) > 20:
                word_count = len(content.split())
                lines.append(f"- {label}: {word_count} words of intelligence gathered")
            else:
                lines.append(f"- {label}: No adverse findings in public domain")
        return "\n".join(lines)

    research_evidence = summarise_web_context(web_context)
    search_date = research_payload.get("research_timestamp", datetime.now().isoformat())[:10]

    return f"""
You are an AI Credit Committee generating a formal Credit Appraisal Memo for an Indian corporate borrower.
Act as FIVE specialist agents in sequence:

AGENT 1 — FINANCIAL ANALYST: Analyze revenue trends, margins, leverage, cashflows.
AGENT 2 — RISK ANALYST: Interpret PD, LGD, Z-Score, fraud flags, triangulation conflicts.
AGENT 3 — INDUSTRY INTELLIGENCE: Assess sector outlook for Indian {sector} sector in 2025.
AGENT 4 — LEGAL & COMPLIANCE RISK: Evaluate legal signals, MCA compliance, e-Courts records,
  CIBIL Commercial, GST filing compliance, RBI adherence.
  MANDATORY: Address GSTR-2A vs 3B reconciliation explicitly with numbers.
AGENT 5 — EWI ANALYST: Identify early warning indicators from all signals including
  triangulation contradictions between web intelligence and document data.

BORROWER DATA:
Company: {company_name} | CIN: {cin} | Sector: {sector}
Loan: ₹{loan_amount}Cr | Decision: {decision} | Max Approved: ₹{max_loan}Cr @ {rate}%
Altman Z-Score: {z_score} — {z_band_narrative}
Note: Do NOT describe this as 'firmly in grey zone' — use the exact narrative above.
PD: {pd_val*100:.1f}% | LGD: {lgd_val*100:.1f}% | Confidence: {confidence*100:.0f}%

FINANCIALS (₹Cr):
Revenue: {financials.get('revenue_cr','N/A')} | EBITDA: {financials.get('ebitda_cr','N/A')}
PAT: {financials.get('pat_cr','N/A')} | Total Debt: {financials.get('total_debt_cr','N/A')}
Net Worth: {financials.get('net_worth_cr','N/A')} | Op. Cashflow: {financials.get('operating_cashflow_cr','N/A')}
Total Assets: {financials.get('total_assets_cr','N/A')}

INDIA-SPECIFIC INDICATORS:
CIBIL Commercial Score: {cibil}/10 | GSTR-2A vs 3B Variance: {gst_var}%
e-Courts Active Cases: {ecourt} | MCA Last Filing: {mca_dt}
RBI Compliance: {rbi_c} | Triangulation Conflicts: {len(tri_flags)}

INDIA-SPECIFIC CONCERNS: {json.dumps(india_concerns)}
SWOT: {swot}
FIVE Cs: {five_cs}
DECISION RATIONALE: {rationale}
FRAUD FLAGS: {flags_s}
TRIANGULATION FLAGS: {tri_s}
EARLY WARNING SIGNALS: {ewi_s}
SECTOR SUMMARY: {ro.get('sector_summary', '')}

WEB INTELLIGENCE CONDUCTED ON {search_date}:
{research_evidence}

INSTRUCTION FOR SECTION 4 (Promoter Background) AND SECTION 11 (Early Warning Indicators):
You MUST include a sentence like: "Web intelligence search conducted on {search_date} [found/found no] adverse public domain information beyond [specific findings or 'the items noted above']."
This is mandatory — do not omit it.

INSTRUCTION FOR SECTION 6 (Industry Outlook):
You MUST reference the specific sector intelligence gathered from web search.
Cite at least one specific policy, regulation, or market development found.

Generate ALL 12 CAM sections in professional Indian banking language.
Be specific. Use exact numbers. Do NOT invent data.
Return valid JSON with these exact keys:
committee_summary, executive_summary, company_overview, promoter_background,
financial_analysis, industry_outlook, legal_risks, five_cs_assessment,
swot_analysis, key_risk_drivers, early_warning_indicators, final_recommendation.
Each value: a string of 2-4 professional paragraphs.
The legal_risks section MUST explicitly address:
- GSTR-2A vs 3B status with actual variance %
- CIBIL Commercial score and implication
- MCA filing currency and days gap
- e-Courts litigation count and nature
- Any triangulation conflicts detected
"""

if "cam_json" not in st.session_state:
    with st.spinner("Multi-agent AI generating 12 CAM sections..."):
        try:
            raw = call_gemini_with_retry([build_cam_prompt()], response_mime_type="application/json")
            st.session_state["cam_json"] = json.loads(raw)
        except Exception as e:
            st.error(f"CAM generation failed: {e}")
            st.session_state["cam_json"] = {s: f"[AI generation failed — please edit manually. Error: {e}]" for s in CAM_SECTIONS}

cam_json: dict = st.session_state["cam_json"]

# ── STEP 3: Section confidence scoring ───────────────────────────────────────
SECTION_KEY_MAP = {
    "financial_analysis":      ["revenue_cr", "ebitda_cr", "total_debt_cr", "operating_cashflow_cr"],
    "legal_risks":             ["cibil_commercial_score", "gst_2a_vs_3b_variance_pct"],
    "industry_outlook":        ["sector_risk_score", "sector_summary"],
    "early_warning_indicators":["early_warning_signals", "triangulation_flags"],
    "five_cs_assessment":      ["five_cs"],
    "swot_analysis":           ["swot"],
    "final_recommendation":    ["decision", "recommended_loan_cr"],
}
LEGAL_EXTRA = {"legal_risk_score": ro.get("legal_risk_score"), "mca_risk_score": ro.get("mca_risk_score")}

def score_section(section: str) -> float:
    base = 0.85 - section_confidence_penalty
    keys = SECTION_KEY_MAP.get(section, [])
    if not keys:
        return round(max(base, 0.40), 2)
    present = sum(1 for k in keys if financials.get(k) or ro.get(k) or REC.get(k))
    completeness = present / len(keys)
    if section == "legal_risks":
        legal_present = sum(1 for v in LEGAL_EXTRA.values() if v)
        completeness = (completeness + legal_present / len(LEGAL_EXTRA)) / 2
    score = base * completeness
    return round(min(max(score, 0.20), 0.97), 2)

confidence_scores = {s: score_section(s) for s in CAM_SECTIONS}

# ── STEP 4: HITL Review interface ─────────────────────────────────────────────
svg_header("EDIT", "Review & Edit CAM Sections", level=2)
st.caption("Each section is AI-generated. Edit as needed — changes are tracked for the audit log.")

web_context = research_payload.get("web_context_used", {})
search_date = research_payload.get("research_timestamp", datetime.now().isoformat())[:10]
st.caption(
    f"🔍 Web intelligence conducted: {search_date} | "
    f"Searches: {len([v for v in web_context.values() if v and len(v) > 20])}/6 returned results | "
    f"Triangulation checks: 5 | Conflicts found: {len(research_payload.get('triangulation_flags', []))}"
)

SECTION_LABELS = {
    "committee_summary":       "1. Committee Summary",
    "executive_summary":       "2. Executive Summary",
    "company_overview":        "3. Company Overview",
    "promoter_background":     "4. Promoter Background",
    "financial_analysis":      "5. Financial Analysis",
    "industry_outlook":        "6. Industry Outlook",
    "legal_risks":             "7. Legal & Compliance Risks",
    "five_cs_assessment":      "8. Five Cs Assessment",
    "swot_analysis":           "9. SWOT Analysis",
    "key_risk_drivers":        "10. Key Risk Drivers",
    "early_warning_indicators":"11. Early Warning Indicators",
    "final_recommendation":    "12. Final Recommendation",
}

for section in CAM_SECTIONS:
    sc = confidence_scores[section]
    conf_label = "High confidence" if sc > 0.80 else ("Medium — review recommended" if sc > 0.60 else "Low — manual input required")
    conf_icon = ":material/check_circle:" if sc > 0.80 else (":material/warning:" if sc > 0.60 else ":material/error:")
    label = SECTION_LABELS.get(section, section.replace("_", " ").title())

    with st.expander(f"{label}  [{sc:.0%} confidence]"):
        if sc > 0.80:
            st.success(conf_label, icon=conf_icon)
        elif sc > 0.60:
            st.warning(conf_label, icon=conf_icon)
        else:
            st.error(conf_label, icon=conf_icon)

        original_text = cam_json.get(section, "")
        edited = st.text_area(
            label=label,
            value=str(original_text),
            height=200,
            key=f"cam_{section}",
            label_visibility="collapsed",
        )
        changed = edited != original_text
        st.session_state[f"cam_original_{section}"] = original_text
        st.session_state[f"cam_approved_{section}"]  = edited
        st.session_state[f"cam_changed_{section}"]   = changed
        if changed:
            st.info("Modified by analyst", icon=":material/edit:")

# ── STEP 5: Financial table + chart ──────────────────────────────────────────
rev   = float(financials.get("revenue_cr")            or 0)
ebit  = float(financials.get("ebitda_cr")             or 0)
pat   = float(financials.get("pat_cr")                or 0)
debt  = float(financials.get("total_debt_cr")         or 0)
nw    = float(financials.get("net_worth_cr")          or 1)
ocf   = float(financials.get("operating_cashflow_cr") or 0)

gst_variance  = (
    financials.get("gst_2a_vs_3b_variance_pct")
    or extraction_payload.get("gst_variance_pct")
    or entity.get("gst_variance_pct")
    or "0.0"   # default — zero variance if not computed
)

mca_date_raw  = entity.get("mca_last_filing_date")
if mca_date_raw:
    try:
        from datetime import datetime, date
        if isinstance(mca_date_raw, str):
            mca_date = datetime.strptime(mca_date_raw[:10], "%Y-%m-%d").date()
        else:
            mca_date = mca_date_raw
        mca_gap_days = (date.today() - mca_date).days
        mca_display  = f"{mca_date_raw} ({mca_gap_days} days ago)"
    except Exception:
        mca_display = str(mca_date_raw)
        mca_gap_days = 0 # fallback
else:
    mca_display = "Not Provided"
    mca_gap_days = 0 

cibil = (
    financials.get("cibil_commercial_score")
    or entity.get("cibil_commercial_score")
    or "N/A"
)

dscr = round(
    (float(financials.get("operating_cashflow_cr") or 0)) /
    max(float(financials.get("total_debt_cr") or 1) * 0.12, 0.01), 2
)

de_ratio = round(
    (float(financials.get("total_debt_cr") or 0)) /
    max(float(financials.get("net_worth_cr") or 1), 0.01), 2
)

# Build the deterministic table — NEVER use LLM for these numbers
financial_summary_rows = [
    ("Revenue (₹ Cr)",              financials.get("revenue_cr", "N/A")),
    ("EBITDA (₹ Cr)",               financials.get("ebitda_cr", "N/A")),
    ("PAT (₹ Cr)",                  financials.get("pat_cr", "N/A")),
    ("Total Debt (₹ Cr)",           financials.get("total_debt_cr", "N/A")),
    ("D/E Ratio",                   f"{de_ratio}x"),
    ("DSCR",                        f"{dscr}x"),
    ("GSTR-2A vs 3B Variance (%)",  f"{gst_variance}%"),
    ("CIBIL Commercial Score",      f"{cibil}/10"),
    ("MCA Last Filing",             mca_display),
]

financial_table = pd.DataFrame(financial_summary_rows, columns=["Metric", "FY2024"])

india_compliance_rows = [
    ("CIBIL Commercial Score",   f"{cibil}/10",      "Low Risk" if float(str(cibil).replace('/10','') or 7) >= 7 else "High Risk"),
    ("GSTR-2A vs 3B Variance",   f"{gst_variance}%", "High Risk" if float(str(gst_variance) or 0) > 15 else "Low Risk"),
    ("MCA Last Filing",          mca_display,         "High Risk" if mca_gap_days > 365 else "Medium Risk" if mca_gap_days > 180 else "Low Risk"),
    ("e-Courts Active Cases",    str(entity.get("ecourt_cases_count", 0)), "High Risk" if int(entity.get("ecourt_cases_count", 0)) > 5 else "Low Risk"),
    ("RBI Compliance",           entity.get("rbi_circular_status", "N/A"), "Low Risk" if "Compliant" in str(entity.get("rbi_circular_status","")) else "High Risk"),
    ("Triangulation Conflicts",  str(len(research_payload.get("triangulation_flags", []))), "High Risk" if len(research_payload.get("triangulation_flags", [])) > 0 else "Low Risk"),
]

def generate_chart() -> str:
    path = "./data/revenue_chart.png"
    cats  = ["Revenue", "EBITDA", "PAT", "Op. Cashflow"]
    vals  = [rev, ebit, pat, ocf]
    clrs  = ["#1B3A6B", "#2563EB", "#059669", "#f59e0b"]
    fig, ax = plt.subplots(figsize=(10, 5))
    bars = ax.bar(cats, vals, color=clrs, width=0.55, edgecolor="white", linewidth=1.5)
    for bar, val in zip(bars, vals):
        ax.text(bar.get_x() + bar.get_width() / 2, bar.get_height() + max(vals)*0.01,
                f"₹{val:.1f}Cr", ha="center", va="bottom", fontweight="bold", fontsize=11)
    ax.set_title(f"Financial Profile — {company_name} (₹ Cr)", fontsize=14, fontweight="bold", color="#1B3A6B")
    ax.set_ylabel("₹ Crore", fontsize=11)
    ax.spines[["top", "right"]].set_visible(False)
    ax.set_facecolor("#f8fafc")
    fig.patch.set_facecolor("white")
    plt.tight_layout()
    plt.savefig(path, dpi=150, bbox_inches="tight")
    plt.close(fig)
    return path

svg_header("CHART", "Financial Summary (Deterministic)", level=2)
st.dataframe(financial_table, use_container_width=True, hide_index=True)

with st.expander("Data Lineage — Extraction Audit Trail", expanded=False):
    lineage_data = extraction_payload.get("data_lineage", {})
    lineage_rows = []
    for k, v in list(lineage_data.items())[:20]:
        val = financials.get(k) or entity.get(k) or "N/A"
        lineage_rows.append({"Field": k, "Extracted Value": val, "Source & Page": v[:100]})
    if lineage_rows:
        st.dataframe(pd.DataFrame(lineage_rows), use_container_width=True, hide_index=True)
    else:
        st.info("No data lineage audit trail available for this scenario.")

# ── STEP 6 & 7: Generate PDFs + Word on button click ─────────────────────────
st.divider()
svg_header("FOLDER", "Export Credit Appraisal Memo", level=2)

approved_sections = {s: st.session_state.get(f"cam_approved_{s}", cam_json.get(s, "")) for s in CAM_SECTIONS}

def build_pdf(chart_path: str) -> bytes:
    from reportlab.lib.pagesizes import A4
    from reportlab.lib import colors
    from reportlab.lib.units import cm
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.platypus import (SimpleDocTemplate, Paragraph, Spacer,
                                    Table, TableStyle, Image, PageBreak, HRFlowable)
    from reportlab.lib.enums import TA_CENTER, TA_LEFT

    buf = io.BytesIO()
    doc = SimpleDocTemplate(buf, pagesize=A4,
                             leftMargin=2*cm, rightMargin=2*cm,
                             topMargin=2.5*cm, bottomMargin=2*cm)

    NAVY   = colors.HexColor("#1B3A6B")
    BLUE   = colors.HexColor("#2563EB")
    GREEN  = colors.HexColor("#059669")
    YELLOW = colors.HexColor("#d97706")
    RED    = colors.HexColor("#dc2626")
    DEC_COLOR = {"APPROVE": GREEN, "MANUAL_REVIEW": YELLOW, "REJECT": RED}.get(decision, NAVY)

    styles = getSampleStyleSheet()
    H1  = ParagraphStyle("h1",  parent=styles["Heading1"],  textColor=NAVY, fontSize=16, spaceAfter=8)
    H2  = ParagraphStyle("h2",  parent=styles["Heading2"],  textColor=NAVY, fontSize=13, spaceAfter=6)
    BOD = ParagraphStyle("bod", parent=styles["Normal"],    fontSize=9.5, leading=14, spaceAfter=6)
    SML = ParagraphStyle("sml", parent=styles["Normal"],    fontSize=8,   leading=12, textColor=colors.grey)
    BIG = ParagraphStyle("big", parent=styles["Title"],     textColor=colors.white,  fontSize=28, alignment=TA_CENTER)
    SUB = ParagraphStyle("sub", parent=styles["Normal"],    textColor=colors.white,  fontSize=13, alignment=TA_CENTER)
    DEC = ParagraphStyle("dec", parent=styles["Title"],     textColor=colors.white,  fontSize=22, alignment=TA_CENTER)

    story = []

    # ── Cover page ─────────────────────────────────────────────────────────
    cover_table = Table([[Paragraph("CREDIT APPRAISAL MEMO", BIG)],
                          [Paragraph("Intelli-Credit AI Platform — CONFIDENTIAL", SUB)]],
                         colWidths=[17*cm])
    cover_table.setStyle(TableStyle([
        ("BACKGROUND", (0,0), (-1,-1), NAVY),
        ("ROWBACKGROUNDS", (0,0), (-1,-1), [NAVY, NAVY]),
        ("TOPPADDING", (0,0), (-1,-1), 30),
        ("BOTTOMPADDING", (0,0), (-1,-1), 30),
        ("LEFTPADDING", (0,0), (-1,-1), 20),
    ]))
    story += [cover_table, Spacer(1, 1*cm)]

    meta = [
        ["Company:", company_name],
        ["CIN:", cin],
        ["Sector:", sector],
        ["Date:", datetime.now().strftime("%d %B %Y")],
        ["Prepared by:", "Intelli-Credit AI Platform v4.0"],
    ]
    meta_tbl = Table(meta, colWidths=[5*cm, 12*cm])
    meta_tbl.setStyle(TableStyle([
        ("FONTNAME", (0,0), (-1,-1), "Helvetica"),
        ("FONTSIZE", (0,0), (-1,-1), 10),
        ("FONTNAME", (0,0), (0,-1), "Helvetica-Bold"),
        ("TEXTCOLOR", (0,0), (0,-1), NAVY),
        ("ROWBACKGROUNDS", (0,0), (-1,-1), [colors.whitesmoke, colors.white]),
        ("GRID", (0,0), (-1,-1), 0.3, colors.lightgrey),
        ("TOPPADDING", (0,0), (-1,-1), 6),
        ("BOTTOMPADDING", (0,0), (-1,-1), 6),
    ]))
    story += [meta_tbl, PageBreak()]

    # ── Page 2: Executive Summary + Decision ────────────────────────────────
    dec_tbl = Table([[Paragraph(f"DECISION: {decision.replace('_',' ')}", DEC)]],
                     colWidths=[17*cm])
    dec_tbl.setStyle(TableStyle([
        ("BACKGROUND", (0,0), (-1,-1), DEC_COLOR),
        ("ROUNDEDCORNERS", [8,8,8,8], 0, 0),
        ("TOPPADDING", (0,0), (-1,-1), 18),
        ("BOTTOMPADDING", (0,0), (-1,-1), 18),
    ]))
    story += [dec_tbl, Spacer(1, 0.5*cm)]

    kpi = [
        ["Recommended Loan", f"Rs {max_loan:.0f} Cr @ {rate:.1f}% p.a."],
        ["PD / LGD", f"{pd_val*100:.1f}% / {lgd_val*100:.1f}%"],
        ["Altman Z-Score", f"{z_score:.2f}"],
        ["Model Confidence", f"{confidence*100:.0f}%"],
    ]
    kpi_tbl = Table(kpi, colWidths=[6*cm, 11*cm])
    kpi_tbl.setStyle(TableStyle([
        ("FONTNAME", (0,0), (-1,-1), "Helvetica"),
        ("FONTNAME", (0,0), (0,-1), "Helvetica-Bold"),
        ("FONTSIZE", (0,0), (-1,-1), 10),
        ("GRID", (0,0), (-1,-1), 0.3, colors.lightgrey),
        ("ROWBACKGROUNDS", (0,0), (-1,-1), [colors.whitesmoke, colors.white]),
        ("TOPPADDING", (0,0), (-1,-1), 6),
    ]))
    story += [Spacer(1, 0.2*cm), kpi_tbl, Spacer(1, 0.5*cm)]

    rationale = REC.get("decision_rationale", [])
    if rationale:
        story.append(Paragraph("Decision Rationale", H2))
        for i, pt in enumerate(rationale, 1):
            story.append(Paragraph(f"{i}. {pt}", BOD))
        story.append(Spacer(1, 0.3*cm))

    if india_concerns:
        story.append(Paragraph("India-Specific Concerns", H2))
        for c in india_concerns:
            story.append(Paragraph(f"• {c}", BOD))

    story.append(PageBreak())

    # ── CAM sections ─────────────────────────────────────────────────────────
    for section in CAM_SECTIONS:
        title = SECTION_LABELS.get(section, section.replace("_", " ").title())
        content = approved_sections.get(section, "")
        story.append(Paragraph(title, H1))
        story.append(HRFlowable(width="100%", thickness=1, color=NAVY))
        story.append(Spacer(1, 0.2*cm))
        for para in content.split("\n\n"):
            if para.strip():
                story.append(Paragraph(para.strip(), BOD))
        story.append(PageBreak())

    # ── India Compliance Summary ───────────────────────────────────────────
    story.append(Paragraph("India-Specific Compliance Summary", H1))
    story.append(HRFlowable(width="100%", thickness=1, color=NAVY))
    story.append(Spacer(1, 0.3*cm))

    india_rows = [["Indicator", "Value", "Risk Band"]]
    for r in india_compliance_rows:
        india_rows.append(list(r))
    comp_tbl = Table(india_rows, colWidths=[6*cm, 5*cm, 6*cm])
    comp_tbl.setStyle(TableStyle([
        ("BACKGROUND", (0,0), (-1,0), NAVY),
        ("TEXTCOLOR", (0,0), (-1,0), colors.white),
        ("FONTNAME", (0,0), (-1,0), "Helvetica-Bold"),
        ("FONTSIZE", (0,0), (-1,-1), 9),
        ("GRID", (0,0), (-1,-1), 0.3, colors.lightgrey),
        ("ROWBACKGROUNDS", (0,1), (-1,-1), [colors.whitesmoke, colors.white]),
        ("TOPPADDING", (0,0), (-1,-1), 6),
        ("BOTTOMPADDING", (0,0), (-1,-1), 6),
    ]))
    story += [comp_tbl, PageBreak()]

    # ── Financial summary ─────────────────────────────────────────────────────
    story.append(Paragraph("Financial Summary", H1))
    story.append(HRFlowable(width="100%", thickness=1, color=NAVY))
    story.append(Spacer(1, 0.3*cm))
    fin_rows = [["Metric", "FY2024"]]
    for _, row in financial_table.iterrows():
        fin_rows.append([str(row["Metric"]), str(row["FY2024"])])
    fin_tbl = Table(fin_rows, colWidths=[9*cm, 8*cm])
    fin_tbl.setStyle(TableStyle([
        ("BACKGROUND", (0,0), (-1,0), NAVY),
        ("TEXTCOLOR", (0,0), (-1,0), colors.white),
        ("FONTNAME", (0,0), (-1,0), "Helvetica-Bold"),
        ("FONTSIZE", (0,0), (-1,-1), 9.5),
        ("GRID", (0,0), (-1,-1), 0.3, colors.lightgrey),
        ("ROWBACKGROUNDS", (0,1), (-1,-1), [colors.whitesmoke, colors.white]),
        ("TOPPADDING", (0,0), (-1,-1), 7),
    ]))
    story += [fin_tbl, Spacer(1, 0.5*cm)]
    if os.path.exists(chart_path):
        story.append(Image(chart_path, width=14*cm, height=7*cm))
    story.append(PageBreak())

    # ── SWOT ─────────────────────────────────────────────────────────────────
    swot_d = REC.get("swot", {})
    if swot_d:
        story.append(Paragraph("SWOT Analysis", H1))
        sw_rows = [
            [Paragraph("<b>Strengths</b>", BOD), Paragraph("<b>Opportunities</b>", BOD)],
            [Paragraph("\n".join(f"• {x}" for x in swot_d.get("strengths",[])), BOD),
             Paragraph("\n".join(f"• {x}" for x in swot_d.get("opportunities",[])), BOD)],
            [Paragraph("<b>Weaknesses</b>", BOD), Paragraph("<b>Threats</b>", BOD)],
            [Paragraph("\n".join(f"• {x}" for x in swot_d.get("weaknesses",[])), BOD),
             Paragraph("\n".join(f"• {x}" for x in swot_d.get("threats",[])), BOD)],
        ]
        sw_tbl = Table(sw_rows, colWidths=[8.5*cm, 8.5*cm])
        sw_tbl.setStyle(TableStyle([
            ("GRID", (0,0), (-1,-1), 0.5, colors.lightgrey),
            ("BACKGROUND", (0,0), (-1,0), colors.HexColor("#dbeafe")),
            ("BACKGROUND", (0,2), (-1,2), colors.HexColor("#fef3c7")),
            ("TOPPADDING", (0,0), (-1,-1), 8),
            ("BOTTOMPADDING", (0,0), (-1,-1), 8),
            ("LEFTPADDING", (0,0), (-1,-1), 8),
        ]))
        story += [sw_tbl, PageBreak()]

    # ── Five Cs ───────────────────────────────────────────────────────────────
    fcs = REC.get("five_cs", {})
    if fcs:
        story.append(Paragraph("Five Cs Assessment", H1))
        fc_rows = [["C", "Score /10", "Comment"]]
        for c_key in ["character","capacity","capital","collateral","conditions"]:
            c_data = fcs.get(c_key, {})
            fc_rows.append([c_key.title(), str(c_data.get("score","N/A")), c_data.get("comment","")])
        fc_tbl = Table(fc_rows, colWidths=[3*cm, 3*cm, 11*cm])
        fc_tbl.setStyle(TableStyle([
            ("BACKGROUND", (0,0), (-1,0), NAVY),
            ("TEXTCOLOR",  (0,0), (-1,0), colors.white),
            ("FONTNAME",   (0,0), (-1,0), "Helvetica-Bold"),
            ("FONTSIZE",   (0,0), (-1,-1), 9.5),
            ("GRID",       (0,0), (-1,-1), 0.3, colors.lightgrey),
            ("ROWBACKGROUNDS", (0,1), (-1,-1), [colors.whitesmoke, colors.white]),
            ("TOPPADDING", (0,0), (-1,-1), 7),
        ]))
        story += [fc_tbl, PageBreak()]

    # ── Appendix: data lineage ────────────────────────────────────────────────
    story.append(PageBreak())
    story.append(Paragraph("Appendix — Data Lineage & Extraction Audit", H1))
    story.append(HRFlowable(width="100%", thickness=1, color=NAVY))
    story.append(Spacer(1, 0.3*cm))
    if lineage_rows:
        lin_rows = [["Field", "Extracted Value", "Source & Page"]]
        for r in lineage_rows:
            lin_rows.append([r["Field"], str(r["Extracted Value"])[:40], str(r["Source & Page"])])
        lin_tbl = Table(lin_rows, colWidths=[4*cm, 4*cm, 9*cm])
        lin_tbl.setStyle(TableStyle([
            ("BACKGROUND", (0,0), (-1,0), NAVY),
            ("TEXTCOLOR", (0,0), (-1,0), colors.white),
            ("FONTNAME", (0,0), (-1,0), "Helvetica-Bold"),
            ("FONTSIZE", (0,0), (-1,-1), 8),
            ("GRID", (0,0), (-1,-1), 0.3, colors.lightgrey),
            ("ROWBACKGROUNDS", (0,1), (-1,-1), [colors.whitesmoke, colors.white]),
            ("VALIGN", (0,0), (-1,-1), "TOP"),
        ]))
        story.append(lin_tbl)
    else:
        story.append(Paragraph("Data lineage not available for this session.", BOD))

    # Confidence scores
    story.append(Spacer(1, 0.5*cm))
    story.append(Paragraph("Section Confidence Scores", H2))
    cs_rows = [["Section", "Confidence"]]
    for s, sc in confidence_scores.items():
        cs_rows.append([SECTION_LABELS.get(s, s), f"{sc:.0%}"])
    cs_tbl = Table(cs_rows, colWidths=[10*cm, 7*cm])
    cs_tbl.setStyle(TableStyle([
        ("FONTSIZE", (0,0), (-1,-1), 9),
        ("GRID", (0,0), (-1,-1), 0.3, colors.lightgrey),
        ("ROWBACKGROUNDS", (0,1), (-1,-1), [colors.whitesmoke, colors.white]),
    ]))
    story.append(cs_tbl)
    story.append(Spacer(1, 0.3*cm))
    story.append(Paragraph(
        f"Generated by: Intelli-Credit AI Platform | Timestamp: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}",
        SML
    ))

    doc.build(story)
    buf.seek(0)
    return buf.read()


def build_word() -> bytes:
    from docx import Document as DocxDoc
    from docx.shared import Pt, RGBColor, Inches, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = DocxDoc()

    # Title
    t = doc.add_heading("CREDIT APPRAISAL MEMO", 0)
    t.runs[0].font.color.rgb = RGBColor(0x1B, 0x3A, 0x6B)
    t.runs[0].font.size = Pt(26)
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Metadata table
    meta = doc.add_table(rows=5, cols=2)
    meta.style = "Table Grid"
    for i, (label, val) in enumerate([
        ("Company", company_name), ("CIN", cin), ("Sector", sector),
        ("Date", datetime.now().strftime("%d %B %Y")),
        ("Prepared by", "Intelli-Credit AI Platform v4.0"),
    ]):
        meta.cell(i, 0).text = label
        meta.cell(i, 1).text = str(val)
    doc.add_paragraph()

    # Decision banner
    dp = doc.add_paragraph()
    dr = dp.add_run(f"DECISION: {decision.replace('_',' ')} — Rs {max_loan:.0f}Cr @ {rate:.1f}%")
    dr.bold, dr.font.size = True, Pt(14)

    # KPI table
    doc.add_heading("Key Risk Metrics", 2)
    kpi_tbl = doc.add_table(rows=4, cols=2)
    kpi_tbl.style = "Table Grid"
    for i, (label, val) in enumerate([
        ("PD / LGD", f"{pd_val*100:.1f}% / {lgd_val*100:.1f}%"),
        ("Altman Z-Score", f"{z_score:.2f}"),
        ("Confidence", f"{confidence*100:.0f}%"),
        ("Max Loan", f"Rs {max_loan:.0f} Cr @ {rate:.1f}%"),
    ]):
        kpi_tbl.cell(i, 0).text = label
        kpi_tbl.cell(i, 1).text = val

    # India-specific concerns
    if india_concerns:
        doc.add_heading("India-Specific Concerns", 2)
        for c in india_concerns:
            doc.add_paragraph(f"• {c}", style="List Bullet")

    doc.add_page_break()

    # 12 CAM sections
    for section in CAM_SECTIONS:
        doc.add_heading(SECTION_LABELS.get(section, section.replace("_", " ").title()), 1)
        content = approved_sections.get(section, "")
        for para in content.split("\n\n"):
            if para.strip():
                doc.add_paragraph(para.strip())
        doc.add_paragraph()

    doc.add_page_break()

    # Financial summary
    doc.add_heading("Financial Summary", 1)
    ft = doc.add_table(rows=len(financial_table)+1, cols=2)
    ft.style = "Table Grid"
    ft.cell(0, 0).text = "Metric"
    ft.cell(0, 1).text = "FY2024"
    for i, row in financial_table.iterrows():
        ft.cell(i+1, 0).text = str(row["Metric"])
        ft.cell(i+1, 1).text = str(row["FY2024"])

    if os.path.exists("./data/revenue_chart.png"):
        doc.add_paragraph()
        doc.add_picture("./data/revenue_chart.png", width=Inches(6))

    doc.add_page_break()

    # India Compliance Summary
    doc.add_heading("India-Specific Compliance Summary", 1)
    ict = doc.add_table(rows=7, cols=3)
    ict.style = "Table Grid"
    headers = ["Indicator", "Value", "Risk Band"]
    for j, h in enumerate(headers):
        ict.cell(0, j).text = h
    for i, (a,b,c) in enumerate(india_compliance_rows, 1):
        ict.cell(i,0).text = a
        ict.cell(i,1).text = str(b)
        ict.cell(i,2).text = str(c)

    doc.add_page_break()

    # SWOT
    swot_d = REC.get("swot", {})
    if swot_d:
        doc.add_heading("SWOT Analysis", 1)
        for quad in ["strengths","weaknesses","opportunities","threats"]:
            doc.add_heading(quad.title(), 3)
            for item in swot_d.get(quad, []):
                doc.add_paragraph(f"• {item}", style="List Bullet")

    # Five Cs
    fcs = REC.get("five_cs", {})
    if fcs:
        doc.add_heading("Five Cs Assessment", 1)
        fc_tbl = doc.add_table(rows=6, cols=3)
        fc_tbl.style = "Table Grid"
        fc_tbl.cell(0,0).text = "C"
        fc_tbl.cell(0,1).text = "Score"
        fc_tbl.cell(0,2).text = "Comment"
        for i, c_key in enumerate(["character","capacity","capital","collateral","conditions"], 1):
            cd = fcs.get(c_key, {})
            fc_tbl.cell(i,0).text = c_key.title()
            fc_tbl.cell(i,1).text = str(cd.get("score","N/A"))
            fc_tbl.cell(i,2).text = cd.get("comment","")

    doc.add_page_break()

    # Appendix: data lineage
    doc.add_heading("Appendix — Data Lineage & Extraction Audit", 1)
    if lineage_rows:
        lin_tbl = doc.add_table(rows=len(lineage_rows)+1, cols=3)
        lin_tbl.style = "Table Grid"
        lin_tbl.cell(0,0).text = "Field"
        lin_tbl.cell(0,1).text = "Extracted Value"
        lin_tbl.cell(0,2).text = "Source & Page"
        for i, r in enumerate(lineage_rows, 1):
            lin_tbl.cell(i,0).text = r["Field"]
            lin_tbl.cell(i,1).text = str(r["Extracted Value"])[:40]
            lin_tbl.cell(i,2).text = str(r["Source & Page"])
    else:
        doc.add_paragraph("Data lineage not available for this session.")

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()


# ── Approve All & generate button ─────────────────────────────────────────────
st.divider()
if st.button("Approve All & Generate CAM", icon=":material/check_circle:", type="primary", use_container_width=True):
    chart_path = generate_chart()
    today_str  = date.today().strftime("%Y%m%d")
    pdf_fname  = f"CAM_{company_name.replace(' ', '_')}_{today_str}.pdf"
    word_fname = f"CAM_{company_name.replace(' ', '_')}_{today_str}.docx"
    pdf_path   = f"./data/{pdf_fname}"

    with st.spinner("Generating PDF..."):
        try:
            pdf_bytes = build_pdf(chart_path)
            with open(pdf_path, "wb") as pf:
                pf.write(pdf_bytes)
        except Exception as ex:
            st.error(f"PDF generation error: {ex}")
            pdf_bytes = None

    with st.spinner("Generating Word document..."):
        try:
            word_bytes = build_word()
        except Exception as ex:
            st.error(f"Word generation error: {ex}")
            word_bytes = None

    # ── STEP 8: Version control ────────────────────────────────────────────
    def get_next_version() -> str:
        try:
            with open("./data/cam_audit_log.json") as fl:
                log = json.load(fl)
            return f"v{len(log) + 1}"
        except (FileNotFoundError, json.JSONDecodeError):
            return "v1"

    version = get_next_version()
    log_entry = {
        "version":   version,
        "timestamp": datetime.now().isoformat(),
        "company":   company_name,
        "decision":  decision,
        "analyst_overrides": {
            s: {
                "original_ai_text":    st.session_state.get(f"cam_original_{s}", ""),
                "analyst_edited_text": st.session_state.get(f"cam_approved_{s}", ""),
                "changed":             st.session_state.get(f"cam_changed_{s}",  False),
            } for s in CAM_SECTIONS
        },
        "confidence_scores": confidence_scores,
        "file_path": pdf_path,
    }

    try:
        with open("./data/cam_audit_log.json") as al:
            audit_log = json.load(al)
    except (FileNotFoundError, json.JSONDecodeError):
        audit_log = []
    audit_log.append(log_entry)
    with open("./data/cam_audit_log.json", "w") as al:
        json.dump(audit_log, al, indent=2, default=str)

    st.success(f"CAM {version} generated for **{company_name}**", icon=":material/verified:")

    # ── Download buttons ───────────────────────────────────────────────────
    col1, col2 = st.columns(2)
    with col1:
        if pdf_bytes:
            st.download_button(
                "Download CAM (PDF)", data=pdf_bytes,
                icon=":material/download:",
                file_name=pdf_fname, mime="application/pdf",
                use_container_width=True,
            )
    with col2:
        if word_bytes:
            st.download_button(
                "Download CAM (Word)", data=word_bytes,
                icon=":material/download:",
                file_name=word_fname,
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                use_container_width=True,
            )

# ── STEP 8: Version history UI ────────────────────────────────────────────────
st.divider()
svg_header("REPORT", "CAM Version History", level=3)
try:
    with open("./data/cam_audit_log.json") as al:
        audit_log = json.load(al)

    if not audit_log:
        st.info("No CAM versions generated yet.")
    else:
        for entry in reversed(audit_log[-5:]):
            overrides_changed = [s for s, v in entry.get("analyst_overrides", {}).items() if v.get("changed")]
            with st.expander(f"**{entry['version']}** — {entry['timestamp'][:16].replace('T',' ')} — {entry['decision']}"):
                st.write(f"**Company:** {entry.get('company','')}")
                st.write(f"**Decision:** {entry.get('decision','')}")
                st.write(f"**File:** `{entry.get('file_path','')}`")
                if overrides_changed:
                    st.write(f"**Analyst-edited sections ({len(overrides_changed)}):** {', '.join(overrides_changed)}")
                else:
                    st.write("**Analyst-edited sections:** None (all AI-generated)")
                cs = entry.get("confidence_scores", {})
                if cs:
                    avg_c = sum(cs.values()) / len(cs)
                    st.write(f"**Avg section confidence:** {avg_c:.0%}")
except (FileNotFoundError, json.JSONDecodeError):
    st.info("No previous versions found.")

# ── Bottom navigation ─────────────────────────────────────────────────────────
st.divider()
n1, n2 = st.columns(2)
with n1:
    if st.button("← Back to Module 3"):
        st.switch_page("pages/03_recommendation.py")
with n2:
    st.success("You have completed all 4 modules of Intelli-Credit.", icon=":material/task_alt:")
